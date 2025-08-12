import os
import time
import shutil
import zipfile
from docx import Document
from docx.shared import Inches
import pandas as pd


def zip_folder(folder_path, output_zip_path):
    """Create a zip file from a folder."""
    with zipfile.ZipFile(
        output_zip_path, "w", zipfile.ZIP_DEFLATED, strict_timestamps=False
    ) as zipf:
        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(
                    file_path, folder_path
                )  # Preserve folder structure
                zipf.write(file_path, arcname)


def create_mapping_dict(df):
    """Create a mapping dictionary from DataFrame."""
    # Take the first row of the data frame and create a mapping dict {df.loc[] : column_name for column_name in df.columns}
    # Select columns where the first row is not empty
    non_empty_columns = df.loc[0].dropna().index
    mapping_dict = {df.loc[0][key].strip(
        " "): key for key in non_empty_columns}
    return mapping_dict


def set_date_and_place(doc):
    """Set date and place in the document."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace("[date]", time.strftime("%d/%m/%Y"))
            run.text = run.text.replace(
                "[date_du_jour]", time.strftime("%d/%m/%Y"))
            run.text = run.text.replace("[Fait_a]", "Arles")


def replace_text(doc, old_text, new_text):
    """Replace text in the document."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(old_text, new_text)


def replace_first_image_in_header(
    doc, new_image_path="logo.png", width_inches=1, height_inches=1
):
    """Replace the first image in the header."""
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if run.element.xpath(".//a:blip"):
                    run.clear()
                    run.add_picture(
                        new_image_path,
                        width=Inches(width_inches),
                        height=Inches(height_inches),
                    )
                    return
