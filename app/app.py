import os
import time
import shutil
import hmac
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from tqdm import tqdm
import concurrent.futures
import gc

from Replacer import WordReplace
from ExcelReplacer import ExcelReplace

import zipfile


def zip_folder(folder_path, output_zip_path):
    with zipfile.ZipFile(
        output_zip_path, "w", zipfile.ZIP_DEFLATED, strict_timestamps=False
    ) as zipf:
        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(
                    file_path, folder_path
                )  # Preserve folder structure
                zipf.write(file_path, arcname)


def create_mapping_dict(df):
    # Take the first row of the data frame and create a mapping dict {df.loc[] : column_name for column_name in df.columns}
    # Select columns where the first row is not empty
    non_empty_columns = df.loc[0].dropna().index
    mapping_dict = {df.loc[0][key].strip(" "): key for key in non_empty_columns}
    return mapping_dict


def set_date_and_place(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace("[date]", time.strftime("%d/%m/%Y"))
            run.text = run.text.replace("[date_du_jour]", time.strftime("%d/%m/%Y"))
            run.text = run.text.replace("[Fait_a]", "Arles")


def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(old_text, new_text)


def replace_first_image_in_header(
    doc, new_image_path="logo.png", width_inches=1, height_inches=1
):
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if run.element.xpath(".//a:blip"):
                    run.clear()
                    run.add_picture(
                        new_image_path,
                        width=Inches(width_inches),
                        height=Inches(height_inches),
                    )
                    return


def check_password():
    """Returns `True` if the user had a correct password."""

    def login_form():
        """Form with widgets to collect user information"""
        with st.form("Credentials"):
            st.text_input("Username", key="username")
            st.text_input("Password", type="password", key="password")
            st.form_submit_button("Log in", on_click=password_entered)

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["username"] in st.secrets[
            "passwords"
        ] and hmac.compare_digest(
            st.session_state["password"],
            st.secrets.passwords[st.session_state["username"]],
        ):
            st.session_state["password_correct"] = True
            # Don't store the username or password.
            del st.session_state["password"]
            del st.session_state["username"]
        else:
            st.session_state["password_correct"] = False

    # Return True if the username + password is validated.
    if st.session_state.get("password_correct", False):
        return True

    # Show inputs for username + password.
    login_form()
    if "password_correct" in st.session_state:
        st.error("😕 User not known or password incorrect")
    return False


def process_word_document(args):
    """
    Process a single Word document - designed for parallel execution
    """
    file_path, mapping_dict, logo_path, template_folder_path, output_folder_path = args
    
    try:
        wordreplace = WordReplace(file_path)
        wordreplace.replace_doc(mapping_dict)
        doc = wordreplace.docx
        set_date_and_place(doc)

        if logo_path and os.path.exists(logo_path):
            replace_first_image_in_header(doc, logo_path)

        doc_name = f"{os.path.basename(file_path)}"
        rel_path = os.path.relpath(file_path, template_folder_path)
        path_to_save = os.path.join(output_folder_path, rel_path)
        path_to_save = path_to_save.replace(os.path.basename(file_path), doc_name)
        doc.save(path_to_save)
        return True, file_path
    except Exception as e:
        return False, f"Error processing {os.path.basename(file_path)}: {str(e)}"


def process_excel_document(args):
    """
    Process a single Excel document - designed for parallel execution
    """
    file_path, mapping_dict, template_folder_path, output_folder_path = args
    
    try:
        excel_replace = ExcelReplace(file_path)
        excel_replace.replace_excel(mapping_dict)
        excel_replace.set_date_and_place()
        
        excel_name = f"{os.path.basename(file_path)}"
        rel_path = os.path.relpath(file_path, template_folder_path)
        path_to_save = os.path.join(output_folder_path, rel_path)
        path_to_save = path_to_save.replace(os.path.basename(file_path), excel_name)
        excel_replace.save(path_to_save)
        return True, file_path
    except Exception as e:
        return False, f"Error processing {os.path.basename(file_path)}: {str(e)}"


# Create the Streamlit app
def main():

    # if not check_password():
    #     st.stop()

    st.title("Générateur de dossier Qualiopi :rocket: ")

    st.sidebar.title("Depot d'excel :newspaper: ")

    excel = st.sidebar.file_uploader(
        "Uploader votre fichier excel", type=["csv", "xlsx", "xls"]
    )

    logo = st.sidebar.file_uploader("Uploader votre logo", type=["png", "jpg", "jpeg"])

    # Performance configuration
    st.sidebar.title("Configuration Performance :gear:")
    use_parallel = st.sidebar.checkbox("Utiliser le traitement parallèle", value=True, help="Active le traitement parallèle pour améliorer les performances")
    max_workers = st.sidebar.slider("Nombre de workers parallèles", min_value=1, max_value=8, value=4, help="Nombre de documents traités simultanément")

    if not excel:
        st.warning("Veuillez uploader un fichier excel pour commencer.")

    if excel:

        if logo is not None:
            # Save the uploaded file to the current directory
            with open("logo.png", "wb") as f:
                f.write(logo.getvalue())

        df = pd.read_excel(excel)
        df = df.astype(str)

        st.sidebar.write(df.iloc[:, 1:].tail(7))

        row_index = st.number_input(
            "Entrer l'indice de ligne prévisualisé à gauche",
            min_value=0,
            max_value=len(df) - 1,
            step=1,
        )

        if st.button("Preview du client"):
            row_data = df.iloc[row_index][
                ["Nom de l'organisme", "Prénom et Nom du responsable de l'organisme"]
            ]
            st.write(row_data)

        if os.path.exists("templates"):
            template_folder_path = "templates"
        else:
            template_folder_path = "app/templates"

        if st.button("Générer les documents") and template_folder_path:
            start_time = time.time()
            
            nom_organisme = df.iloc[row_index]["Nom de l'organisme"]
            # Create a folder to store generated documents
            output_folder_path = f"docs/{nom_organisme}_{time.strftime('%H_%M_%S')}"
            # Copy template structure (now including Excel files)
            shutil.copytree(
                template_folder_path,
                output_folder_path,
                ignore=shutil.ignore_patterns("*.docx"),
            )

            mappings = create_mapping_dict(df)

            mapping_dict = {
                key: str(df.iloc[row_index][value]) for key, value in mappings.items()
            }

            progress_bar = st.progress(0, text=f"Progress: 0%")
            
            # Process both Word and Excel documents
            doc_list = WordReplace.docx_list(template_folder_path)
            excel_list = ExcelReplace.excel_list(template_folder_path)
            total_files = len(doc_list) + len(excel_list)
            
            st.info(f"Traitement de {len(doc_list)} documents Word et {len(excel_list)} documents Excel")
            
            file_counter = 0

            # Process Word documents
            if use_parallel and len(doc_list) > 1:
                with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = [
                        executor.submit(process_word_document, (file, mapping_dict, "logo.png" if logo is not None else None, template_folder_path, output_folder_path))
                        for file in doc_list
                    ]
                    for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures)):
                        file_counter += 1
                        progress_bar.progress(
                            file_counter / total_files,
                            text=f"Document Word {file_counter}/{total_files}",
                        )
                        success, result = future.result()
                        if not success:
                            st.warning(result)  # result contains the error message
                        # Force garbage collection to free memory
                        gc.collect()
            else:
                # Sequential processing for small document sets or when parallel is disabled
                for i, file in tqdm(enumerate(doc_list)):
                    file_counter += 1
                    progress_bar.progress(
                        file_counter / total_files,
                        text=f"Document Word {file_counter}/{total_files}",
                    )
                    success, result = process_word_document((file, mapping_dict, "logo.png" if logo is not None else None, template_folder_path, output_folder_path))
                    if not success:
                        st.warning(result)
                    gc.collect()

            # Process Excel documents (new feature)
            if use_parallel and len(excel_list) > 1:
                with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = [
                        executor.submit(process_excel_document, (file, mapping_dict, template_folder_path, output_folder_path))
                        for file in excel_list
                    ]
                    for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures)):
                        file_counter += 1
                        progress_bar.progress(
                            file_counter / total_files,
                            text=f"Document Excel {file_counter}/{total_files}",
                        )
                        success, result = future.result()
                        if not success:
                            st.warning(result)  # result contains the error message
                        # Force garbage collection to free memory
                        gc.collect()
            else:
                # Sequential processing for small document sets or when parallel is disabled
                for i, file in tqdm(enumerate(excel_list)):
                    file_counter += 1
                    progress_bar.progress(
                        file_counter / total_files,
                        text=f"Document Excel {file_counter}/{total_files}",
                    )
                    success, result = process_excel_document((file, mapping_dict, template_folder_path, output_folder_path))
                    if not success:
                        st.warning(result)
                    gc.collect()

            zip_folder(output_folder_path, output_folder_path + ".zip")

            end_time = time.time()
            processing_time = end_time - start_time
            
            st.success(f"Documents générés en {processing_time:.2f} secondes ! Dossier: {output_folder_path}")
            st.info(f"Performance: {total_files/processing_time:.2f} documents/seconde")

            with open(output_folder_path + ".zip", "rb") as f:
                st.download_button(
                    label="Télécharger le dossier des documents générés",
                    data=f,
                    file_name=output_folder_path + ".zip",
                    mime="application/zip",
                )

        if st.button("Supprimer le dossier généré"):
            if os.path.exists("docs"):
                shutil.rmtree("docs")
                # print("Folder deleted")


if __name__ == "__main__":
    main()
