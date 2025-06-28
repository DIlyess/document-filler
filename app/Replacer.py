#!/usr/bin/env python 3.9
# -*- coding: utf-8 -*-
# @Time    : 2022/12/6 18:02 update
# @Author  : ZCG
# @File    : WordReplace.py
# @Software: PyCharm
# @Notice  : Optimized version for better performance with large document sets


from docx import Document
import os
import re


class OptimizedExecute:
    """
    Optimized Execute Paragraphs KeyWords Replace
    paragraph: docx paragraph
    """

    def __init__(self, paragraph):
        self.paragraph = paragraph

    def replace_all_in_paragraph(self, replace_dict: dict):
        """
        Optimized paragraph replacement - processes all replacements in one pass
        """
        if not self.paragraph.text or not replace_dict:
            return

        # Get all runs and their text
        runs = list(self.paragraph.runs)
        if not runs:
            return

        # Create a single text string and track run boundaries
        full_text = ""
        run_boundaries = []
        
        for run in runs:
            start_pos = len(full_text)
            full_text += run.text
            end_pos = len(full_text)
            run_boundaries.append((start_pos, end_pos, run))

        # Apply all replacements to the full text
        modified_text = full_text
        replacements = []
        
        for key, value in replace_dict.items():
            if key in modified_text:
                # Find all occurrences
                start = 0
                while True:
                    pos = modified_text.find(key, start)
                    if pos == -1:
                        break
                    replacements.append((pos, pos + len(key), value))
                    start = pos + 1

        if not replacements:
            return

        # Sort replacements by position (reverse order to avoid index shifting)
        replacements.sort(key=lambda x: x[0], reverse=True)

        # Apply replacements
        for start_pos, end_pos, replacement in replacements:
            modified_text = modified_text[:start_pos] + replacement + modified_text[end_pos:]

        # Reconstruct runs with the modified text
        self._reconstruct_runs(run_boundaries, modified_text)

    def _reconstruct_runs(self, run_boundaries, new_text):
        """
        Reconstruct runs with new text while preserving formatting
        """
        if not run_boundaries:
            return

        # Clear all runs except the first one
        first_run = run_boundaries[0][2]
        for _, _, run in run_boundaries[1:]:
            run.text = ""

        # Calculate new run boundaries based on original proportions
        total_original_length = sum(end - start for start, end, _ in run_boundaries)
        if total_original_length == 0:
            return

        # Distribute new text across runs proportionally
        current_pos = 0
        for i, (start, end, run) in enumerate(run_boundaries):
            original_length = end - start
            proportion = original_length / total_original_length
            new_length = int(len(new_text) * proportion)
            
            if i == len(run_boundaries) - 1:
                # Last run gets remaining text
                run.text = new_text[current_pos:]
            else:
                run.text = new_text[current_pos:current_pos + new_length]
                current_pos += new_length


class OptimizedWordReplace:
    """
    Optimized Word document processing for better performance
    file: Microsoft Office word file，only support .docx type file
    """

    def __init__(self, file):
        self.docx = Document(file)
        self._cached_sections = None
        self._cached_tables = None

    def _get_sections(self):
        """Cache sections to avoid repeated access"""
        if self._cached_sections is None:
            self._cached_sections = list(self.docx.sections)
        return self._cached_sections

    def _get_tables(self):
        """Cache tables to avoid repeated access"""
        if self._cached_tables is None:
            self._cached_tables = list(self.docx.tables)
        return self._cached_tables

    def _process_paragraphs(self, paragraphs, replace_dict):
        """Process a collection of paragraphs efficiently"""
        if not replace_dict:
            return
            
        for paragraph in paragraphs:
            if paragraph.text and any(key in paragraph.text for key in replace_dict.keys()):
                OptimizedExecute(paragraph).replace_all_in_paragraph(replace_dict)

    def body_content(self, replace_dict: dict):
        """Optimized body content replacement"""
        self._process_paragraphs(self.docx.paragraphs, replace_dict)

    def body_tables(self, replace_dict: dict):
        """Optimized body tables replacement"""
        if not replace_dict:
            return
            
        for table in self._get_tables():
            for row in table.rows:
                for cell in row.cells:
                    self._process_paragraphs(cell.paragraphs, replace_dict)

    def header_content(self, replace_dict: dict):
        """Optimized header content replacement"""
        if not replace_dict:
            return
            
        for section in self._get_sections():
            if section.header:
                self._process_paragraphs(section.header.paragraphs, replace_dict)

    def header_tables(self, replace_dict: dict):
        """Optimized header tables replacement"""
        if not replace_dict:
            return
            
        for section in self._get_sections():
            if section.header:
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            self._process_paragraphs(cell.paragraphs, replace_dict)

    def footer_content(self, replace_dict: dict):
        """Optimized footer content replacement"""
        if not replace_dict:
            return
            
        for section in self._get_sections():
            if section.footer:
                self._process_paragraphs(section.footer.paragraphs, replace_dict)

    def footer_tables(self, replace_dict: dict):
        """Optimized footer tables replacement"""
        if not replace_dict:
            return
            
        for section in self._get_sections():
            if section.footer:
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            self._process_paragraphs(cell.paragraphs, replace_dict)

    def save(self, filepath: str):
        """Save the modified document"""
        self.docx.save(filepath)

    @staticmethod
    def docx_list(dirPath):
        """Get list of docx files in directory and subdirectories"""
        file_list = []
        for roots, dirs, files in os.walk(dirPath):
            for file in files:
                # Find the docx document and exclude temporary files
                if file.endswith("docx") and file[0] != "~":
                    file_root = os.path.join(roots, file)
                    file_list.append(file_root)
        return file_list

    def replace_doc(self, replace_dict: dict):
        """Optimized document replacement - processes all sections efficiently"""
        if not replace_dict:
            return self.docx

        # Process all content types in one pass
        self.header_content(replace_dict)
        self.body_content(replace_dict)
        self.footer_content(replace_dict)
        # Uncomment if table processing is needed
        # self.body_tables(replace_dict)
        # self.header_tables(replace_dict)
        # self.footer_tables(replace_dict)

        return self.docx


# Backward compatibility - keep the old class names
class Execute(OptimizedExecute):
    """Backward compatibility wrapper"""
    pass


class WordReplace(OptimizedWordReplace):
    """Backward compatibility wrapper"""
    pass


def main():
    """
    Example usage
    """
    # Example replace dictionary
    replace_dict = {
        "aaa": "bbb",
        "ccc": "ddd",
    }
    filedir = r"D:\Working Files\svn"

    # Process all files
    for i, file in enumerate(WordReplace.docx_list(filedir), start=1):
        print(f"{i}. Processing file: {file}")
        wordreplace = WordReplace(file)
        wordreplace.replace_doc(replace_dict)
        wordreplace.save(file)
        print(f"Document processing complete!")


if __name__ == "__main__":
    main()
